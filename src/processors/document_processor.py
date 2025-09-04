"""
Document processing system for handling company documents and bid materials.
"""
import os
import re
from typing import List, Dict, Any, Optional, Union
from pathlib import Path
from dataclasses import dataclass
from loguru import logger

import PyPDF2
from docx import Document
import pandas as pd
from openpyxl import load_workbook

@dataclass
class ProcessedDocument:
    """Data class representing a processed document."""
    filename: str
    file_type: str
    content: str
    metadata: Dict[str, Any]
    extracted_keywords: List[str]
    sections: Dict[str, str]

class DocumentProcessor:
    """Processes various document types for bid applications."""
    
    def __init__(self, documents_folder: str = "./documents"):
        self.documents_folder = Path(documents_folder)
        self.documents_folder.mkdir(exist_ok=True)
        
        # Common keywords to extract from documents
        self.technical_keywords = [
            "cybersecurity", "information security", "IT services", "software development",
            "network security", "cloud computing", "data protection", "risk assessment",
            "compliance", "penetration testing", "vulnerability assessment", "incident response",
            "security operations center", "SOC", "SIEM", "firewall", "encryption",
            "authentication", "authorization", "access control", "monitoring", "logging"
        ]
        
        self.certification_keywords = [
            "certified", "certification", "ISO", "SOC", "CMMI", "ITIL", "PMP",
            "CISSP", "CISM", "CISA", "CEH", "Security+", "CISSP", "CISM"
        ]
        
        self.experience_keywords = [
            "experience", "years", "implemented", "managed", "developed", "designed",
            "architected", "deployed", "maintained", "supported", "delivered"
        ]
    
    def process_all_documents(self) -> List[ProcessedDocument]:
        """Process all documents in the documents folder."""
        processed_docs = []
        
        if not self.documents_folder.exists():
            logger.warning(f"Documents folder {self.documents_folder} does not exist")
            return processed_docs
        
        # Process all supported file types
        for file_path in self.documents_folder.rglob("*"):
            if file_path.is_file() and self._is_supported_file(file_path):
                try:
                    doc = self.process_document(file_path)
                    if doc:
                        processed_docs.append(doc)
                        logger.info(f"Processed document: {file_path.name}")
                except Exception as e:
                    logger.error(f"Failed to process {file_path}: {e}")
        
        logger.info(f"Processed {len(processed_docs)} documents")
        return processed_docs
    
    def process_document(self, file_path: Union[str, Path]) -> Optional[ProcessedDocument]:
        """Process a single document."""
        file_path = Path(file_path)
        
        if not file_path.exists():
            logger.error(f"File does not exist: {file_path}")
            return None
        
        file_type = self._get_file_type(file_path)
        content = self._extract_content(file_path, file_type)
        
        if not content:
            logger.warning(f"No content extracted from {file_path}")
            return None
        
        # Extract metadata and keywords
        metadata = self._extract_metadata(file_path, content)
        keywords = self._extract_keywords(content)
        sections = self._extract_sections(content)
        
        return ProcessedDocument(
            filename=file_path.name,
            file_type=file_type,
            content=content,
            metadata=metadata,
            extracted_keywords=keywords,
            sections=sections
        )
    
    def _is_supported_file(self, file_path: Path) -> bool:
        """Check if file type is supported."""
        supported_extensions = {'.pdf', '.docx', '.doc', '.txt', '.xlsx', '.xls', '.csv'}
        return file_path.suffix.lower() in supported_extensions
    
    def _get_file_type(self, file_path: Path) -> str:
        """Get file type from extension."""
        return file_path.suffix.lower().lstrip('.')
    
    def _extract_content(self, file_path: Path, file_type: str) -> str:
        """Extract content from different file types."""
        try:
            if file_type == 'pdf':
                return self._extract_pdf_content(file_path)
            elif file_type in ['docx', 'doc']:
                return self._extract_docx_content(file_path)
            elif file_type == 'txt':
                return self._extract_txt_content(file_path)
            elif file_type in ['xlsx', 'xls']:
                return self._extract_excel_content(file_path)
            elif file_type == 'csv':
                return self._extract_csv_content(file_path)
            else:
                logger.warning(f"Unsupported file type: {file_type}")
                return ""
        except Exception as e:
            logger.error(f"Failed to extract content from {file_path}: {e}")
            return ""
    
    def _extract_pdf_content(self, file_path: Path) -> str:
        """Extract text content from PDF."""
        content = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    content += page.extract_text() + "\n"
        except Exception as e:
            logger.error(f"Failed to extract PDF content: {e}")
        return content
    
    def _extract_docx_content(self, file_path: Path) -> str:
        """Extract text content from DOCX."""
        content = ""
        try:
            doc = Document(file_path)
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"
        except Exception as e:
            logger.error(f"Failed to extract DOCX content: {e}")
        return content
    
    def _extract_txt_content(self, file_path: Path) -> str:
        """Extract text content from TXT."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read()
            except Exception as e:
                logger.error(f"Failed to extract TXT content: {e}")
                return ""
    
    def _extract_excel_content(self, file_path: Path) -> str:
        """Extract text content from Excel files."""
        content = ""
        try:
            workbook = load_workbook(file_path, data_only=True)
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                content += f"Sheet: {sheet_name}\n"
                for row in sheet.iter_rows(values_only=True):
                    row_text = " ".join(str(cell) for cell in row if cell is not None)
                    if row_text.strip():
                        content += row_text + "\n"
        except Exception as e:
            logger.error(f"Failed to extract Excel content: {e}")
        return content
    
    def _extract_csv_content(self, file_path: Path) -> str:
        """Extract text content from CSV files."""
        content = ""
        try:
            df = pd.read_csv(file_path)
            content = df.to_string()
        except Exception as e:
            logger.error(f"Failed to extract CSV content: {e}")
        return content
    
    def _extract_metadata(self, file_path: Path, content: str) -> Dict[str, Any]:
        """Extract metadata from document."""
        metadata = {
            'file_size': file_path.stat().st_size,
            'word_count': len(content.split()),
            'character_count': len(content),
            'line_count': len(content.split('\n'))
        }
        
        # Extract company information patterns
        company_patterns = [
            r'(?i)Company\s+Profile\s*:\s*([^\n]+)',
            r'(?i)Company\s*Name\s*:\s*([^\n]+)',
            r'(?i)Company\s*:\s*([^\n]+)',
            r'(?i)Organization\s*:\s*([^\n]+)',
            r'(?i)Firm\s*:\s*([^\n]+)'
        ]
        
        for pattern in company_patterns:
            match = re.search(pattern, content, re.IGNORECASE)
            if match:
                name = match.group(1).strip()
                # Clean up cases like "Profile: TechSecure Solutions"
                if name.lower().startswith('profile:'):
                    name = name.split(':', 1)[-1].strip()
                metadata['company_name'] = name
                break
        
        # Extract contact information
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, content)
        if emails:
            metadata['emails'] = emails
        
        phone_pattern = r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
        phones = re.findall(phone_pattern, content)
        if phones:
            metadata['phones'] = phones
        
        return metadata
    
    def _extract_keywords(self, content: str) -> List[str]:
        """Extract relevant keywords from content."""
        content_lower = content.lower()
        found_keywords = []
        
        # Check for technical keywords
        for keyword in self.technical_keywords:
            if keyword.lower() in content_lower:
                found_keywords.append(keyword)
        
        # Check for certification keywords
        for keyword in self.certification_keywords:
            if keyword.lower() in content_lower:
                found_keywords.append(keyword)
        
        # Check for experience keywords
        for keyword in self.experience_keywords:
            if keyword.lower() in content_lower:
                found_keywords.append(keyword)
        
        return list(set(found_keywords))  # Remove duplicates
    
    def _extract_sections(self, content: str) -> Dict[str, str]:
        """Extract document sections based on common headings."""
        sections = {}
        
        # Common section patterns
        section_patterns = {
            'executive_summary': r'(?i)(executive\s+summary|summary)',
            'company_overview': r'(?i)(company\s+overview|about\s+us|organization)',
            'technical_capabilities': r'(?i)(technical\s+capabilities|capabilities|services)',
            'experience': r'(?i)(experience|past\s+performance|project\s+history)',
            'certifications': r'(?i)(certifications|certificates|credentials)',
            'team': r'(?i)(team|personnel|staff|key\s+personnel)',
            'methodology': r'(?i)(methodology|approach|process)'
        }
        
        lines = content.split('\n')
        current_section = None
        current_content = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check if line matches a section header
            section_found = False
            for section_name, pattern in section_patterns.items():
                if re.search(pattern, line):
                    # Save previous section
                    if current_section and current_content:
                        sections[current_section] = '\n'.join(current_content)
                    
                    # Start new section
                    current_section = section_name
                    current_content = [line]
                    section_found = True
                    break
            
            if not section_found and current_section:
                current_content.append(line)
        
        # Save last section
        if current_section and current_content:
            sections[current_section] = '\n'.join(current_content)
        
        return sections
    
    def get_company_profile(self, processed_docs: List[ProcessedDocument]) -> Dict[str, Any]:
        """Create a comprehensive company profile from processed documents."""
        profile = {
            'company_name': '',
            'capabilities': [],
            'certifications': [],
            'experience_areas': [],
            'key_personnel': [],
            'contact_info': {},
            'technical_keywords': set(),
            'all_content': ''
        }
        
        for doc in processed_docs:
            # Combine all content
            profile['all_content'] += doc.content + '\n'
            
            # Extract company name
            if 'company_name' in doc.metadata:
                fname = doc.filename.lower()
                if (not profile['company_name']) or ('company_profile' in fname):
                    profile['company_name'] = doc.metadata['company_name']
            
            # Collect technical keywords
            profile['technical_keywords'].update(doc.extracted_keywords)
            
            # Extract sections
            if 'company_overview' in doc.sections:
                profile['capabilities'].append(doc.sections['company_overview'])
            
            if 'technical_capabilities' in doc.sections:
                profile['capabilities'].append(doc.sections['technical_capabilities'])
            
            if 'certifications' in doc.sections:
                profile['certifications'].append(doc.sections['certifications'])
            
            if 'experience' in doc.sections:
                profile['experience_areas'].append(doc.sections['experience'])
            
            if 'team' in doc.sections:
                profile['key_personnel'].append(doc.sections['team'])
            
            # Extract contact information
            if 'emails' in doc.metadata:
                profile['contact_info']['emails'] = doc.metadata['emails']
            if 'phones' in doc.metadata:
                profile['contact_info']['phones'] = doc.metadata['phones']
        
        # Convert set to list
        profile['technical_keywords'] = list(profile['technical_keywords'])
        
        return profile

